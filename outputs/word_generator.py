"""
Word Document Generator — produces a professionally structured .docx proposal.

Structure:
  - Cover page
  - Table of contents (manual)
  - Technical proposal sections (all 12 sections)
  - Commercial proposal sections
  - Costing table

Uses python-docx styles for Protiviti brand consistency.
The Word output is designed for submission and further editing.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config import OUTPUTS_PATH, BLENDED_RATE_USD

# Protiviti brand colors (as RGB tuples)
RED = RGBColor(0xC8, 0x10, 0x2E)
DARK = RGBColor(0x40, 0x40, 0x40)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT = RGBColor(0xF5, 0xF5, 0xF5)


def generate_word(proposal_data: dict, client_name: str) -> str:
    """Generate Word document. Returns file path."""
    doc = Document()
    _setup_document(doc)

    tech = proposal_data.get("technical", {})
    commercial = proposal_data.get("commercial", {})
    effort = proposal_data.get("effort_model", {})
    rfp = proposal_data.get("rfp_intel", {})
    web = proposal_data.get("web_research", {})
    value_add = proposal_data.get("value_add_slide", {})

    # Cover
    _add_cover(doc, tech.get("cover", {}), client_name)
    doc.add_page_break()

    # Confidentiality notice
    p = doc.add_paragraph("STRICTLY CONFIDENTIAL — Prepared exclusively for " + client_name)
    p.runs[0].font.color.rgb = RED
    p.runs[0].font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # ── Technical Proposal ────────────────────────────────────────────────────
    _section_heading(doc, "TECHNICAL PROPOSAL")

    _write_exec_summary(doc, tech.get("executive_summary", {}))
    _write_client_overview(doc, web.get("client_profile", {}), rfp)
    _write_sector_context(doc, web.get("sector_context", {}), rfp)
    _write_understanding(doc, tech.get("our_understanding", {}))
    _write_value_proposition(doc, tech.get("value_proposition", {}))
    _write_past_relationship(doc, tech.get("past_relationship", {}))
    _write_scope(doc, tech.get("scope_of_work", {}))
    _write_value_add(doc, value_add)
    _write_approach(doc, tech.get("approach_methodology", {}))
    _write_governance(doc, tech.get("engagement_governance", {}))
    _write_team(doc, tech.get("project_team", {}))
    _write_timeline(doc, tech.get("timeline", {}))
    _write_experience(doc, tech.get("relevant_experience", {}))
    _write_why_protiviti(doc, tech.get("why_protiviti", {}))
    _write_assumptions(doc, tech.get("key_assumptions", {}))

    doc.add_page_break()

    # ── Commercial Proposal ───────────────────────────────────────────────────
    _section_heading(doc, "COMMERCIAL PROPOSAL")
    _write_commercial(doc, commercial, effort)

    # Footer
    doc.add_page_break()
    p = doc.add_paragraph(
        "This proposal is confidential and prepared exclusively for "
        + client_name
        + ". Protiviti Middle East — Real Estate & Infrastructure Practice."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    filename = f"Proposal_{client_name.replace(' ', '_')}.docx"
    path = os.path.join(OUTPUTS_PATH, filename)
    doc.save(path)
    return path


# ── Setup ─────────────────────────────────────────────────────────────────────

def _setup_document(doc: Document):
    """Set page margins and default font."""
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default normal style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = DARK

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RED

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = DARK

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(10)
    h3.font.bold = True
    h3.font.color.rgb = GRAY


# ── Helpers ───────────────────────────────────────────────────────────────────

def _section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RED
    _set_paragraph_border_bottom(p)
    doc.add_paragraph("")


def _h1(doc, text):
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(16)
    return h


def _h2(doc, text):
    return doc.add_heading(text, level=2)


def _h3(doc, text):
    return doc.add_heading(text, level=3)


def _body(doc, text, italic=False):
    if not text:
        return
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    if italic:
        run.font.italic = True


def _bullet(doc, text, bold_prefix: str = ""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.font.bold = True
        run.font.size = Pt(10)
    run = p.add_run(str(text))
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY


def _label(doc, text):
    """Red bold label/subheading."""
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.bold = True
    run.font.color.rgb = RED
    run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def _set_paragraph_border_bottom(paragraph):
    """Add a red bottom border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C8102E")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_table(doc, headers: list, rows: list[list]):
    """Add a simple styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        # Red background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "C8102E")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            row.cells[c_idx].text = str(cell_text)
            row.cells[c_idx].paragraphs[0].runs[0].font.size = Pt(9)
            if r_idx % 2 == 0:
                tc = row.cells[c_idx]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F5F5F5")
                shd.set(qn("w:val"), "clear")
                tcPr.append(shd)

    doc.add_paragraph("")


# ── Section writers ───────────────────────────────────────────────────────────

def _add_cover(doc, cover: dict, client_name: str):
    doc.add_paragraph("")
    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("PROTIVITI MIDDLE EAST")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RED
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph()
    run2 = p2.add_run("Real Estate & Infrastructure Practice")
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    doc.add_paragraph("")

    p3 = doc.add_paragraph()
    run3 = p3.add_run(cover.get("title", "Technical Proposal"))
    run3.font.size = Pt(22)
    run3.font.bold = True
    run3.font.color.rgb = DARK
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p4 = doc.add_paragraph()
    run4 = p4.add_run(cover.get("subtitle", "Technical Proposal | Confidential"))
    run4.font.size = Pt(12)
    run4.font.color.rgb = GRAY
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    doc.add_paragraph("")

    p5 = doc.add_paragraph()
    run5 = p5.add_run(f"Prepared for: {cover.get('client', client_name)}")
    run5.font.size = Pt(11)
    run5.font.bold = True
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p6 = doc.add_paragraph()
    run6 = p6.add_run(f"Date: {cover.get('date', 'April 2026')}")
    run6.font.size = Pt(10)
    run6.font.color.rgb = GRAY
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _write_exec_summary(doc, data: dict):
    if not data:
        return
    _h1(doc, "Executive Summary")
    _label(doc, "Our Understanding")
    _body(doc, data.get("our_understanding", ""))
    _label(doc, "Our Approach")
    _body(doc, data.get("our_approach", ""))
    _label(doc, "Our Commitment")
    _body(doc, data.get("our_commitment", ""))
    for o in data.get("key_outcomes", []):
        _bullet(doc, o)


def _write_client_overview(doc, profile: dict, rfp: dict):
    if not profile or profile.get("error"):
        return
    _h1(doc, f"About {rfp.get('client_name', 'the Client')}")
    _body(doc, profile.get("organization_overview", ""))
    _label(doc, "Core Mandate")
    _body(doc, profile.get("mandate_and_role", ""))
    _label(doc, "Recent Initiatives")
    for i in profile.get("recent_initiatives", [])[:4]:
        item = i.get("initiative", i) if isinstance(i, dict) else i
        _bullet(doc, item)
    _label(doc, "Leadership Priorities")
    for p in profile.get("leadership_priorities", [])[:4]:
        _bullet(doc, p)


def _write_sector_context(doc, sector: dict, rfp: dict):
    if not sector or sector.get("error"):
        return
    _h1(doc, f"{rfp.get('sector', 'Sector')} Landscape — {rfp.get('geography', 'UAE')}")
    _body(doc, sector.get("sector_overview", ""))
    _label(doc, "Transformation Drivers")
    for d in sector.get("transformation_drivers", [])[:4]:
        _bullet(doc, d)
    _label(doc, "National Programs")
    for p in sector.get("national_programs", [])[:4]:
        item = p.get("program", p) if isinstance(p, dict) else p
        _bullet(doc, item)
    _label(doc, "Key Challenges")
    for c in sector.get("sector_challenges", [])[:4]:
        _bullet(doc, c)


def _write_understanding(doc, data: dict):
    if not data:
        return
    _h1(doc, data.get("title", "Our Understanding"))
    _body(doc, data.get("context_narrative", ""))
    _label(doc, "Key Drivers")
    for d in data.get("key_drivers", []):
        if isinstance(d, dict):
            _bullet(doc, d.get("detail", ""), bold_prefix=d.get("driver", "") + ":")
        else:
            _bullet(doc, d)
    _label(doc, "Challenges Identified")
    for c in data.get("challenges_identified", []):
        if isinstance(c, dict):
            _bullet(doc, c.get("implication", ""), bold_prefix=c.get("challenge", "") + ":")
        else:
            _bullet(doc, c)
    _label(doc, "Success Factors")
    for f in data.get("success_factors", []):
        if isinstance(f, dict):
            _bullet(doc, f.get("why", ""), bold_prefix=f.get("factor", "") + ":")
        else:
            _bullet(doc, f)
    if data.get("our_perspective"):
        _label(doc, "Our Perspective")
        _body(doc, data.get("our_perspective", ""), italic=True)


def _write_value_proposition(doc, data: dict):
    if not data:
        return
    _h1(doc, data.get("title", "Our Value Proposition"))
    p = doc.add_paragraph()
    run = p.add_run(data.get("headline", ""))
    run.font.bold = True
    run.font.color.rgb = RED
    run.font.size = Pt(11)

    for vp in data.get("value_points", []):
        _label(doc, vp.get("point", ""))
        _body(doc, vp.get("detail", ""))
        if vp.get("proof_point"):
            _body(doc, f"✓ {vp['proof_point']}", italic=True)

    _label(doc, "Our Differentiators")
    for d in data.get("protiviti_differentiators", []):
        _bullet(doc, d)
    if data.get("what_sets_us_apart"):
        _label(doc, "What Sets Us Apart")
        _body(doc, data.get("what_sets_us_apart", ""))


def _write_past_relationship(doc, data: dict):
    if not data or not data.get("relationship_narrative"):
        return
    _h1(doc, data.get("title", "Our Relationship with the Client"))
    _body(doc, data.get("relationship_narrative", ""))
    raw_eng = data.get("past_engagements", [])
    if raw_eng:
        _label(doc, "Past Engagements")
        for e in raw_eng:
            if isinstance(e, dict):
                _bullet(doc, e.get("outcome", ""), bold_prefix=f"{e.get('engagement','')} ({e.get('year','')}):")
            else:
                _bullet(doc, e)
    if data.get("continuity_benefit"):
        _label(doc, "Continuity Benefit")
        _body(doc, data.get("continuity_benefit", ""))


def _write_scope(doc, scope: dict):
    if not scope:
        return
    _h1(doc, scope.get("title", "Scope of Work & Deliverables"))
    _body(doc, scope.get("scope_narrative", ""))
    for phase in scope.get("phases", []):
        _h2(doc, phase.get("phase_name", ""))
        _body(doc, phase.get("phase_objective", ""))
        for d in phase.get("deliverables_l1", []):
            _h3(doc, f"■ {d.get('name', '')}")
            _body(doc, d.get("description", ""))
            for sub in d.get("sub_deliverables", []):
                _bullet(doc, sub)
        if phase.get("phase_milestone"):
            _label(doc, "Phase Milestone")
            _body(doc, phase.get("phase_milestone", ""), italic=True)


def _write_value_add(doc, data: dict):
    if not data:
        return
    _h1(doc, data.get("title", "Our Recommended Enhancements"))
    _body(doc, data.get("intro", ""))
    if data.get("included_in_scope"):
        _label(doc, "Included in Scope — Leading Practice Additions")
        for item in data.get("included_in_scope", []):
            _bullet(doc, item.get("talking_point", item.get("benefit", "")),
                    bold_prefix=f"✓ {item.get('title', '')}:")
    if data.get("optional_addons"):
        _label(doc, "Optional Enhancements")
        for item in data.get("optional_addons", []):
            fee = item.get("fee_usd", 0)
            _bullet(doc, f"{item.get('benefit', '')} — USD {fee:,.0f}",
                    bold_prefix=f"+ {item.get('title', '')}:")


def _write_approach(doc, data: dict):
    if not data:
        return
    _h1(doc, data.get("title", "Our Approach & Methodology"))
    _body(doc, data.get("approach_narrative", ""))
    for p in data.get("guiding_principles", []):
        if isinstance(p, dict):
            _label(doc, p.get("principle", ""))
            _body(doc, p.get("description", ""))
        else:
            _bullet(doc, p)
    _label(doc, "Methodology Steps")
    for i, step in enumerate(data.get("methodology_steps", []), 1):
        _h3(doc, f"{i}. {step.get('step', '')}")
        _body(doc, step.get("description", ""))
        for act in step.get("activities", []):
            _bullet(doc, act)
    _label(doc, "Tools & Accelerators")
    for t in data.get("tools_and_accelerators", []):
        if isinstance(t, dict):
            _bullet(doc, t.get("use_case", ""), bold_prefix=f"{t.get('tool', '')}:")
        else:
            _bullet(doc, t)
    if data.get("knowledge_transfer"):
        _label(doc, "Knowledge Transfer")
        _body(doc, data.get("knowledge_transfer", ""))


def _write_governance(doc, data: dict):
    if not data:
        return
    _h1(doc, data.get("title", "Engagement Governance"))
    _body(doc, data.get("governance_narrative", ""))
    gov = data.get("governance_structure", "")
    if isinstance(gov, dict):
        for k, v in gov.items():
            _label(doc, k.replace("_", " ").title())
            _body(doc, v)
    elif gov:
        _body(doc, gov)
    # Reporting cadence table
    cadence = data.get("reporting_cadence", [])
    if cadence and isinstance(cadence[0], dict):
        _label(doc, "Reporting Cadence")
        _add_table(doc,
                   ["Report", "Frequency", "Audience", "Content"],
                   [[c.get("report",""), c.get("frequency",""),
                     c.get("audience",""), c.get("content","")] for c in cadence])
    # RACI table
    raci = data.get("raci_summary", [])
    if raci:
        _label(doc, "RACI Summary")
        _add_table(doc, ["Activity", "Protiviti", "Client"],
                   [[r.get("activity",""), r.get("protiviti",""), r.get("client","")]
                    for r in raci[:6]])
    if data.get("quality_assurance"):
        _label(doc, "Quality Assurance")
        _body(doc, data.get("quality_assurance", ""))


def _write_team(doc, data: dict):
    if not data:
        return
    _h1(doc, data.get("title", "Our Project Team"))
    _body(doc, data.get("team_narrative", ""))
    for m in data.get("team_members", []):
        _h2(doc, f"{m.get('role', '')} — {m.get('title', '')}")
        if m.get("name"):
            p = doc.add_paragraph()
            run = p.add_run(m["name"])
            run.font.bold = True
        _body(doc, m.get("relevant_experience", ""))
        for r in m.get("responsibilities", []):
            _bullet(doc, r)


def _write_timeline(doc, data: dict):
    if not data:
        return
    _h1(doc, data.get("title", "Engagement Timeline"))
    p = doc.add_paragraph()
    run = p.add_run(f"Total Duration: {data.get('total_duration', '')}")
    run.font.bold = True
    run.font.color.rgb = RED
    _body(doc, data.get("timeline_narrative", ""))
    for phase in data.get("phases", []):
        _h2(doc, f"{phase.get('phase', '')}  |  {phase.get('duration', '')}")
        _label(doc, "Key Activities")
        for act in phase.get("key_activities", []):
            _bullet(doc, act)
        if phase.get("key_milestones"):
            _label(doc, "Milestones")
            for m in phase.get("key_milestones", []):
                _bullet(doc, m)
        if phase.get("client_inputs_required"):
            _label(doc, "Client Inputs Required")
            for ci in phase.get("client_inputs_required", []):
                _bullet(doc, ci)


def _write_experience(doc, data: dict):
    if not data:
        return
    _h1(doc, data.get("title", "Our Relevant Experience"))
    _body(doc, data.get("narrative", ""))
    _body(doc, data.get("regional_credentials", ""))
    for cs in data.get("case_studies", []):
        _h2(doc, f"{cs.get('engagement', '')} — {cs.get('client_type', '')}")
        _body(doc, cs.get("context", ""))
        _label(doc, "Our Role")
        _body(doc, cs.get("our_role", ""))
        _label(doc, "Outcome")
        _body(doc, cs.get("outcome", ""), italic=True)
        _label(doc, "Relevance to This Proposal")
        _body(doc, cs.get("relevance", ""), italic=True)


def _write_why_protiviti(doc, data: dict):
    if not data:
        return
    _h1(doc, data.get("title", "Why Protiviti"))
    p = doc.add_paragraph()
    run = p.add_run(data.get("headline", ""))
    run.font.bold = True
    run.font.color.rgb = RED
    _body(doc, data.get("closing_narrative", ""))
    for r in data.get("reasons", []):
        if isinstance(r, dict):
            _label(doc, f"◆ {r.get('reason', '')}")
            _body(doc, r.get("detail", ""))
        else:
            _bullet(doc, r)


def _write_assumptions(doc, data: dict):
    if not data:
        return
    _h1(doc, data.get("title", "Key Assumptions & Dependencies"))
    _label(doc, "Key Assumptions")
    for a in data.get("assumptions", []):
        _bullet(doc, a)
    _label(doc, "Client Dependencies")
    for d in data.get("client_dependencies", []):
        _bullet(doc, d)
    _label(doc, "Out of Scope")
    for o in data.get("out_of_scope", []):
        _bullet(doc, o)
    if data.get("variation_trigger"):
        _label(doc, "Scope Variation Trigger")
        _body(doc, data.get("variation_trigger", ""))


def _write_commercial(doc, commercial: dict, effort: dict):
    _h1(doc, "Commercial Summary")
    summary = commercial.get("commercial_summary", {})
    total_fee = effort.get("total_fee_usd", summary.get("total_fee_usd", 0))
    total_hours = effort.get("total_hours", summary.get("total_hours", 0))

    _add_table(doc,
               ["Item", "Value"],
               [
                   ["Total Engagement Fee", f"USD {total_fee:,.0f}"],
                   ["Total Hours", f"{total_hours:,}"],
                   ["Blended Rate", f"USD {BLENDED_RATE_USD}/hour"],
                   ["Payment Structure", summary.get("payment_structure", "Milestone-based")],
                   ["Proposal Validity", f"{summary.get('validity_days', 30)} days"],
               ])

    # Fee by phase
    _h2(doc, "Fee by Phase")
    phases = effort.get("phases", commercial.get("fee_by_phase", []))
    if phases:
        rows = []
        for ph in phases:
            name = ph.get("phase_name", ph.get("phase", ""))
            fee = ph.get("phase_fee_usd", ph.get("fee_usd", 0))
            hours = ph.get("phase_hours", ph.get("hours", 0))
            pct = round(fee / total_fee * 100) if total_fee else 0
            rows.append([name, f"{hours:,}", f"USD {fee:,.0f}", f"{pct}%"])
        _add_table(doc, ["Phase", "Hours", "Fee (USD)", "% of Total"], rows)

    # Payment milestones
    milestones = commercial.get("payment_milestones", [])
    if milestones:
        _h2(doc, "Payment Milestones")
        rows = [[
            m.get("milestone", ""), m.get("trigger", ""),
            f"USD {m.get('amount_usd', 0):,.0f}", m.get("percentage", ""),
        ] for m in milestones]
        _add_table(doc, ["Milestone", "Payment Trigger", "Amount", "%"], rows)

    # Terms
    terms = commercial.get("terms", {})
    if terms:
        _h2(doc, "Commercial Terms")
        for k, v in terms.items():
            _label(doc, k.replace("_", " ").title())
            _body(doc, v)
